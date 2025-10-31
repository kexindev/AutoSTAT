import ast
import io
import re
from io import BytesIO

import streamlit as st
from stqdm import stqdm
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.express as px
import plotly.io as pio

from utils.sanitize_code import sanitize_code



def write_word(agents):
    '''
    choice：是否要搜索
        True：根据目录搜索相关章节
        False：全部章节
    '''
    # 拿图
    analysis_list = agents[2].summary_fig_analysis_list()

    report_agent = agents[-1]
    report_obj = report_agent.load_report()  # Reportcore

    doc = Document()

    style = doc.styles['Normal']

    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _insert_figure(fig_obj):
        if fig_obj is None:
            return
        try:
            img_bytes = io.BytesIO()
            img_bytes = io.BytesIO(fig_obj.to_image(format="png"))
            # pio.write_image(fig_obj, img_bytes, format='png')
            img_bytes.seek(0)

            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(img_bytes, width=Inches(4))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[图像插入失败: {e}]")

    def _process_node(node):
        if node.type == "heading":
            doc.add_heading(node.text, level=node.level)
            for ch in node.children:
                _process_node(ch)

        elif node.type == "paragraph":
            parts = re.split(r'(\[FIG:\d+\])', node.text)

            for part in parts:
                part = part.strip()
                if not part:
                    continue
                if part.startswith("[FIG:") and part.endswith("]"):
                    idx = int(part[5:-1])
                    fig_obj = None
                    if 0 <= idx < len(analysis_list):
                        entry = analysis_list[idx]
                        fig_obj = entry.get("figure")
                    _insert_figure(fig_obj)
                else:
                    doc.add_paragraph(part)

        else:  # root
            for ch in node.children:
                _process_node(ch)

    # 从 root.children 开始写
    _process_node(report_obj.root)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    report_agent.save_word(buf.getvalue())
    st.success("Word 报告生成成功")
